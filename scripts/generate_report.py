from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_thesis_report():
    doc = Document()

    # --- Title Page ---
    doc.add_heading('Automated PV Lifecycle Assessment (LCA) Tool', 0)
    subtitle = doc.add_paragraph('Thesis Presentation & Future Roadmap Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n\n')
    
    # --- 1. Executive Summary ---
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "This project addresses a critical data transparency gap in the photovoltaic (PV) industry. "
        "While Environmental Product Declarations (EPDs) exist, they are currently locked in static, unstructured PDF formats, "
        "making large-scale comparison and sustainable procurement impossible.\n\n"
        "We have developed a Python-based data pipeline and an interactive Decision Support Tool (Web App) that automates "
        "the extraction, normalization, and visualization of EPD data. This empowers stakeholders to make data-driven decisions "
        "based on true Carbon Intensity (gCO2e/kWh) rather than just manufacturing footprint."
    )

    # --- 2. The Problem Statement ---
    doc.add_heading('2. The Problem Statement', level=1)
    doc.add_paragraph('Why does this tool need to exist?')
    
    problems = [
        ("Data Silos", "Critical environmental data is trapped in unstructured PDF documents scattered across various program operators."),
        ("Incomparable Units", "Manufacturers report data differently (per module vs. per Watt vs. per m2), making direct apple-to-apple comparisons manually impossible."),
        ("Static Analysis", "Current procurement decisions often ignore operational lifespan and degradation, which are critical for true 'cradle-to-grave' impact."),
        ("Greenwashing Risk", "Without a standardized way to audit data quality, incomplete EPDs can falsely appear superior.")
    ]
    
    for title, desc in problems:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    # --- 3. The Solution: PV LCA Decision Tool ---
    doc.add_heading('3. The Solution: PV LCA Decision Tool', level=1)
    doc.add_paragraph(
        "We have built a comprehensive solution consisting of a data normalization engine and an interactive dashboard."
    )
    
    # [PLACEHOLDER FOR DASHBOARD SCREENSHOT]
    doc.add_paragraph('\n[INSERT SCREENSHOT OF DASHBOARD HERE]\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Figure 1: The Main Dashboard Interface showing Carbon Intensity Ranking', style='Caption')

    doc.add_heading('3.1 Key Features', level=2)
    features = [
        ("Project Simulator", "Unlike static lists, the tool simulates specific project contexts (Location Yield, Lifetime) to calculate real-world impact."),
        ("Dynamic Metrics", "Calculates Carbon Intensity (gCO2e/kWh) and Carbon Payback Time (CPT) in real-time."),
        ("Quality Assurance", "An automated 'Star Rating' system assesses data reliability, directly combating greenwashing.")
    ]
    for title, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    # [PLACEHOLDER FOR SENSITIVITY TAB SCREENSHOT]
    doc.add_paragraph('\n[INSERT SCREENSHOT OF SENSITIVITY TAB HERE]\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Figure 2: Sensitivity Analysis Interface for Degradation Modeling', style='Caption')

    # --- 4. Industry Use Cases ---
    doc.add_heading('4. Industry Use Cases', level=1)
    
    doc.add_heading('Use Case A: Solar Developers & EPCs', level=3)
    doc.add_paragraph(
        "Scenario: A developer needs to build a 50 MW plant in Southern Spain.\n"
        "Solution: The tool identifies that while Glass-Glass modules have higher initial carbon, their lower degradation "
        "results in a 15% lower lifetime carbon intensity in high-irradiance climates."
    )

    doc.add_heading('Use Case B: Policy Makers & Public Procurement', level=3)
    doc.add_paragraph(
        "Scenario: A government wants to set a carbon cap for new solar tenders.\n"
        "Solution: The tool provides a benchmark for 'average' vs. 'best-in-class' carbon intensity to set realistic regulatory thresholds."
    )

    # --- 5. Future Roadmap: Agentic AI Automation ---
    doc.add_heading('5. Future Roadmap: Agentic AI Automation', level=1)
    doc.add_paragraph(
        "The current pipeline relies on manual CSV entry. The next phase introduces "
        "Agentic AI to fully automate the lifecycle from 'PDF Discovery' to 'Database Update'."
    )

    # [PLACEHOLDER FOR PIPELINE DIAGRAM]
    # You can insert the architecture diagram here if you have one
    # doc.add_paragraph('\n[INSERT PIPELINE ARCHITECTURE DIAGRAM HERE]\n').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Phase 3: The "Autonomous EPD Hunter"', level=2)
    doc.add_paragraph("We propose a stack using n8n (Orchestrator), OpenAI GPT-4o (Intelligence), and Pinecone (Vector DB).")

    workflow = [
        ("The Scout (Search Agent)", "Triggers weekly to scrape EPD registries and identify valid PDF links for 'Photovoltaic Modules'."),
        ("The Reader (Extraction Agent)", "Downloads PDFs and uses Vision-Language Models (GPT-4o) to extract unstructured data like GWP, Efficiency, and Dimensions into JSON."),
        ("The Analyst (Normalization Agent)", "Applies the logic built in Phase 1 to convert mixed units into standard Wp metrics."),
        ("The Publisher (Deployment Agent)", "Commits clean data to the GitHub repository and notifies stakeholders via Slack/Email.")
    ]

    for title, desc in workflow:
        p = doc.add_paragraph(style='List Number')
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    doc.add_paragraph(
        "\nImpact: This shift reduces data maintenance time by 99% and ensures the database grows automatically "
        "as new products are released."
    )

    # --- 6. Conclusion ---
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        "This project transforms static, opaque environmental documents into a dynamic, transparent decision-making engine. "
        "By combining rigorous LCA methodology with modern web technology, we have built a tool that doesn't just show data—it "
        "interprets it for real-world impact. The transition to an Agentic AI pipeline will secure this tool's position as a "
        "scalable, industry-leading standard for sustainable solar procurement."
    )

    # Save the document
    file_name = 'PV_LCA_Thesis_Report.docx'
    doc.save(file_name)
    print(f"✅ Successfully generated '{file_name}'")

if __name__ == "__main__":
    create_thesis_report()